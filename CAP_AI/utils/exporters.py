"""Export utilities — Excel, CSV, PDF, Word."""

from __future__ import annotations

import io
from datetime import datetime

import pandas as pd
from fpdf import FPDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def to_excel_bytes(df: pd.DataFrame, sheet_name: str = "Data") -> bytes:
    """Export DataFrame to Excel bytes."""
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return buf.getvalue()


def to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def to_pdf_report(title: str, sections: list[tuple[str, str]], df: pd.DataFrame | None = None) -> bytes:
    """Generate PDF report with ReportLab."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Heading1"], textColor=colors.HexColor("#0ea5e9"))
    body = [
        Paragraph(title, title_style),
        Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]),
        Spacer(1, 0.3 * inch),
    ]
    for heading, content in sections:
        body.append(Paragraph(heading, styles["Heading2"]))
        body.append(Paragraph(content, styles["Normal"]))
        body.append(Spacer(1, 0.2 * inch))

    if df is not None and not df.empty:
        data = [df.columns.tolist()] + df.head(50).values.tolist()
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0ea5e9")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
        ]))
        body.append(table)

    doc.build(body)
    return buf.getvalue()


def to_word_report(title: str, sections: list[tuple[str, str]]) -> bytes:
    """Generate Word document."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def simple_pdf_table(title: str, df: pd.DataFrame) -> bytes:
    """Quick PDF via FPDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, title, ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 8, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
    pdf.ln(5)
    cols = df.columns.tolist()
    pdf.set_font("Helvetica", "B", 8)
    for c in cols:
        pdf.cell(40, 7, str(c)[:15], border=1)
    pdf.ln()
    pdf.set_font("Helvetica", "", 7)
    for _, row in df.head(30).iterrows():
        for c in cols:
            pdf.cell(40, 6, str(row[c])[:15], border=1)
        pdf.ln()
    return pdf.output()
